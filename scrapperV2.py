import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
import time
from collections import deque
import re
from datetime import datetime, timezone
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.shared import qn, OxmlElement
from docx.enum.text import WD_COLOR_INDEX

# Configuration
DEV_BASE_URL = "https://developer.weweb.io"
DOCS_BASE_URL = "https://docs.weweb.io"
HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
}
OUTPUT_DIR = "weweb_docs"
REQUEST_DELAY = 1.2  # Délai entre les requêtes

def sanitize_filename(filename):
    filename = re.sub(r'[\\/*?:"<>|]', '_', filename)
    return filename[:50].strip()

def add_hyperlink(paragraph, url, text):
    """Ajoute un lien cliquable dans un paragraphe Word."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000EE')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def get_all_urls(base_url):
    visited = set()
    queue = deque([base_url])

    while queue:
        url = queue.popleft()
        if url in visited:
            continue

        try:
            print(f"🔍 Exploration de: {url}")
            response = requests.get(url, headers=HEADERS, timeout=15)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')
            visited.add(url)

            for a in soup.find_all('a', href=True):
                href = a['href']
                if not href or href.startswith(('javascript:', 'mailto:', 'tel:', '#')):
                    continue

                absolute_url = urljoin(base_url, href)
                if absolute_url.startswith(base_url):
                    clean_url = absolute_url.split('#')[0].rstrip('/')
                    if clean_url not in visited and clean_url not in queue:
                        queue.append(clean_url)

            time.sleep(REQUEST_DELAY)

        except Exception as e:
            print(f"⚠️ Erreur avec {url}: {str(e)}")
            continue

    return sorted(visited)

def create_docx(page_data, output_dir):
    try:
        doc = Document()

        # Style de base
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        doc.add_heading(page_data['title'], 0)

        # Table des métadonnées
        meta_table = doc.add_table(rows=2, cols=2)
        meta_table.style = 'LightShading-Accent1'

        meta_table.cell(0, 0).text = "URL source"
        add_hyperlink(meta_table.cell(0, 1).paragraphs[0], page_data['url'], page_data['url'])

        meta_table.cell(1, 0).text = "Date de scraping"
        meta_table.cell(1, 1).text = page_data['metadata']['scraped_at']

        for section_id, section in page_data['sections'].items():
            doc.add_heading(section['title'], level=2)

            if section['content']:
                doc.add_paragraph(section['content'])

            if section['code_snippets']:
                doc.add_heading("Code", level=3)
                for snippet in section['code_snippets']:
                    doc.add_paragraph(f"Langage: {snippet['language']}", style='Intense Quote')
                    code_para = doc.add_paragraph()
                    run = code_para.add_run(snippet['code'])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

            if section['images']:
                doc.add_heading("Images", level=3)
                for img_url in section['images']:
                    p = doc.add_paragraph()
                    add_hyperlink(p, img_url, img_url)

        filename = f"{sanitize_filename(page_data['title'])}.docx"
        output_path = os.path.join(output_dir, filename)

        counter = 1
        while os.path.exists(output_path):
            output_path = os.path.join(output_dir, f"{sanitize_filename(page_data['title'])}_{counter}.docx")
            counter += 1

        doc.save(output_path)
        print(f"✅ Fichier créé: {os.path.basename(output_path)}")

    except Exception as e:
        print(f"❌ Erreur DOCX: {str(e)}")
        raise

def scrape_page(url):
    try:
        print(f"⏳ Scraping de {url}")
        response = requests.get(url, headers=HEADERS, timeout=15)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')

        page_data = {
            'url': url,
            'title': soup.title.get_text().strip() if soup.title else url.split('/')[-1],
            'sections': {},
            'metadata': {
                'scraped_at': datetime.now(timezone.utc).isoformat(),
                'source_url': url
            }
        }

        for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5']):
            section_id = heading.get('id')
            if not section_id:
                continue

            section = {
                'title': heading.get_text().strip(),
                'content': [],
                'code_snippets': [],
                'images': []
            }

            next_node = heading.find_next_sibling()
            while next_node and next_node.name not in ['h1', 'h2', 'h3', 'h4', 'h5']:
                if next_node.name:
                    if next_node.name in ['p', 'ul', 'ol']:
                        section['content'].append(next_node.get_text(' ', strip=True))

                    for img in next_node.find_all('img'):
                        src = img.get('src') or img.get('data-src')
                        if src:
                            if not src.startswith(('http://', 'https://')):
                                src = urljoin(url, src)
                            if src not in section['images']:
                                section['images'].append(src)

                    for pre in next_node.find_all('pre'):
                        code = pre.find('code')
                        if code:
                            lang_classes = [cls for cls in code.get('class', []) if cls.startswith('language-')]
                            lang = lang_classes[0].replace('language-', '') if lang_classes else 'unknown'
                            section['code_snippets'].append({
                                'code': code.get_text().strip(),
                                'language': lang
                            })

                next_node = next_node.find_next_sibling()

            section['content'] = '\n'.join(section['content']).strip()
            page_data['sections'][section_id] = section

        return page_data

    except Exception as e:
        print(f"❌ Erreur de scraping: {str(e)}")
        return None

def main():
    print("🚀 Démarrage du scraping WeWeb")
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("\n🔍 Exploration des sites...")
    dev_urls = get_all_urls(DEV_BASE_URL)
    docs_urls = get_all_urls(DOCS_BASE_URL)

    total = len(dev_urls) + len(docs_urls)
    start_time = time.time()

    for idx, url in enumerate(dev_urls + docs_urls, 1):
        print(f"\n📄 Traitement {idx}/{total}: {url}")
        page_data = scrape_page(url)
        if page_data:
            create_docx(page_data, OUTPUT_DIR)
        time.sleep(REQUEST_DELAY)

    print(f"\n✅ Terminé en {time.time() - start_time:.2f} secondes")
    print(f"📂 Dossier de sortie: {os.path.abspath(OUTPUT_DIR)}")

if __name__ == "__main__":
    main()
