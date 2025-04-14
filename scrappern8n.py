import asyncio
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

BASE_URL = "https://docs.n8n.io"
OUTPUT_DIR = "n8n_docs_clean"

os.makedirs(OUTPUT_DIR, exist_ok=True)

async def scrape_and_format_docs():
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page()
        await page.goto(BASE_URL)

        await page.wait_for_selector("nav a")

        # Récupère tous les liens internes du menu
        links = await page.eval_on_selector_all("nav a", "elements => elements.map(e => e.href)")
        links = list(set([link for link in links if link.startswith(BASE_URL)]))

        print(f"🔗 {len(links)} liens trouvés dans le menu")

        for idx, link in enumerate(links):
            try:
                await page.goto(link)
                await page.wait_for_selector("main")
                
                # Sélection du conteneur de contenu principal (à adapter si nécessaire)
                content = await page.query_selector('main .md-content__inner')
                if not content:
                    content = await page.query_selector('main article')  # Fallback

                title = await page.title()
                title = title.replace(" | n8n Docs", "").strip()

                doc = Document()
                doc.add_heading(title, level=0)

                if content:
                    # Sélection des éléments uniquement dans le contenu principal
                    elements = await content.query_selector_all("h1, h2, h3, p, li, pre, code, table")
                    
                    for el in elements:
                        tag = await el.evaluate("e => e.tagName.toLowerCase()")
                        text = (await el.inner_text()).strip()

                        if not text:
                            continue

                        # Gestion des titres
                        if tag in ["h1", "h2", "h3"]:
                            level = int(tag[1])
                            doc.add_heading(text, level=level)
                        
                        # Gestion des paragraphes
                        elif tag == "p":
                            doc.add_paragraph(text)
                        
                        # Gestion des listes
                        elif tag == "li":
                            doc.add_paragraph(f"• {text}")
                        
                        # Gestion du code
                        elif tag in ["pre", "code"]:
                            para = doc.add_paragraph()
                            run = para.add_run(text)
                            run.font.name = "Courier New"
                            run.font.size = Pt(9)
                        
                        # Gestion des tableaux
                        elif tag == "table":
                            rows = await el.query_selector_all("tr")
                            table = doc.add_table(rows=1, cols=1)
                            
                            for row in rows:
                                cells = await row.query_selector_all("td, th")
                                row_data = [await cell.inner_text() for cell in cells]
                                
                                if not table.rows:
                                    table.add_row()
                                row_cells = table.rows[-1].cells
                                
                                for i, data in enumerate(row_data):
                                    if i >= len(row_cells):
                                        table.add_column()
                                        row_cells = table.rows[-1].cells
                                    row_cells[i].text = data

                filename = f"{OUTPUT_DIR}/{idx+1:03d}_{title[:50].replace(' ', '_').replace('/', '-')}.docx"
                doc.save(filename)
                print(f"✅ {filename} sauvegardé")

            except Exception as e:
                print(f"❌ Erreur sur {link}: {str(e)}")

        await browser.close()

asyncio.run(scrape_and_format_docs())